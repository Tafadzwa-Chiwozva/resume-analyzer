from flask import Flask, request, jsonify, send_from_directory
import os
from parser import extract_text_from_pdf, analyze_resume
from pdf2docx import Converter
from docx import Document
from docx2pdf import convert

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
PROCESSED_FOLDER = "processed_resumes"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)

@app.route("/")
def home():
    return "Resume Analyzer Backend is Running!"

@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "Empty filename"}), 400
    
    file_path = os.path.join(UPLOAD_FOLDER, file.filename)
    file.save(file_path)

    # Extract text from resume
    try:
        extracted_text = extract_text_from_pdf(file_path)
        job_category = request.form.get('job_category', 'General')
        ai_feedback = analyze_resume(extracted_text, job_category)
        
        if "optimized_resume" in ai_feedback:
            optimized_text = ai_feedback["optimized_resume"]
            optimized_filename = f"optimized_{file.filename}"
            optimized_file_path = os.path.join(PROCESSED_FOLDER, optimized_filename)
            
            # Convert PDF to DOCX
            docx_path = file_path.replace('.pdf', '.docx')
            cv = Converter(file_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()
            
            # Modify DOCX with optimized text
            modify_docx(docx_path, optimized_text)
            
            # Convert DOCX back to PDF
            final_pdf_path = optimized_file_path
            convert_docx_to_pdf(docx_path, final_pdf_path)
            
            return jsonify({
                "filename": file.filename,
                "extracted_text": extracted_text,
                "ai_feedback": ai_feedback,
                "download_url": f"/download/{optimized_filename}"
            })
        else:
            return jsonify({"error": "Failed to optimize resume"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/download/<filename>", methods=["GET"])
def download_file(filename):
    return send_from_directory(PROCESSED_FOLDER, filename)

def modify_docx(docx_path, optimized_text):
    doc = Document(docx_path)
    # Clear the existing content
    for paragraph in doc.paragraphs:
        paragraph.clear()
    # Add the optimized text
    doc.add_paragraph(optimized_text)
    doc.save(docx_path)

def convert_docx_to_pdf(docx_path, pdf_path):
    convert(docx_path, pdf_path)

if __name__ == "__main__":
    app.run(debug=True)